from typing import Optional, Dict
from file_utils import create_backup
from bible_references import process_text  # Import the function to process text
from docx import Document
import os

def process_paragraph(paragraph) -> bool:
    """
    Process a single paragraph to standardize Bible references.
    Returns True if the paragraph was modified, False otherwise.
    """
    original_text = paragraph.text
    standardized_text = process_text(original_text)
    if original_text != standardized_text:
        paragraph.text = standardized_text
        return True
    return False

def process_document(doc_path: str, output_path: Optional[str] = None, should_create_backup: bool = True) -> Dict:
    """
    Process a Word document to standardize Bible references.
    """
    result = {
        'success': False,
        'changes_made': False,
        'paragraphs_processed': 0,
        'paragraphs_changed': 0,
        'backup_path': None,
        'output_path': output_path or doc_path,
        'error': None
    }

    # Check if input file exists
    if not os.path.exists(doc_path):
        result['error'] = f"File not found: {doc_path}"
        return result

    # Create backup if requested
    if should_create_backup:
        backup_success, backup_result = create_backup(doc_path)
        if backup_success:
            result['backup_path'] = backup_result
        else:
            result['error'] = f"Failed to create backup: {backup_result}"
            return result

    try:
        # Open document
        document = Document(doc_path)

        # Process paragraphs
        for paragraph in document.paragraphs:
            result['paragraphs_processed'] += 1
            if process_paragraph(paragraph):
                result['paragraphs_changed'] += 1
                result['changes_made'] = True

        # Process tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        result['paragraphs_processed'] += 1
                        if process_paragraph(paragraph):
                            result['paragraphs_changed'] += 1
                            result['changes_made'] = True

        # Save document
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        document.save(output_path)
        result['success'] = True

    except Exception as e:
        result['error'] = f"Unexpected error processing document: {str(e)}"

    return result