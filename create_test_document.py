"""
create_test_document.py - Script to create a test document with non-standardized Bible references.
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_test_document(output_path="test_non_standard.docx"):
    """Create a test document with various non-standardized Bible references."""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Test Document with Non-Standardized Bible References', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add paragraphs with various non-standardized references
    doc.add_paragraph('This document contains multiple Bible references in non-standardized formats to test the scripture standardization script.')
    
    # Section 1: Abbreviated book names
    doc.add_heading('Section 1: Abbreviated Book Names', level=1)
    doc.add_paragraph('The apostle Paul encourages believers to put on Christ as the new man (Gal. 3:27*, Eph. 4:24*, Col. 3:10*).')
    doc.add_paragraph('In Rom. 1:16*, Paul declares that he is not ashamed of the gospel.')
    
    # Section 2: Period-separated verses
    doc.add_heading('Section 2: Period-Separated Verses', level=1)
    doc.add_paragraph('In 1 Cor. 13.4-7, Paul describes the characteristics of love.')
    doc.add_paragraph('The prophet Isaiah wrote in Isa. 53.5-6 about the suffering servant.')
    
    # Section 3: Space-separated verses
    doc.add_heading('Section 3: Space-Separated Verses', level=1)
    doc.add_paragraph('Jesus said in Matt 5 3-10 that the meek shall inherit the earth.')
    doc.add_paragraph('Psalm 23 is a well-known passage about God\'s shepherding care.')
    
    # Section 4: Invalid range formats
    doc.add_heading('Section 4: Invalid Range Formats', level=1)
    doc.add_paragraph('Acts 2:12:16 records the reaction of the crowd on the day of Pentecost.')
    doc.add_paragraph('Genesis 1:1:3 describes the first days of creation.')
    
    # Section 5: Chapter-verse format
    doc.add_heading('Section 5: Chapter-Verse Format', level=1)
    doc.add_paragraph('In John chapter 3 verse 16, we learn about God\'s love for the world.')
    doc.add_paragraph('Luke chapter 2 verses 1-7 tells the story of Jesus\' birth.')
    
    # Section 6: Mixed references in parentheses
    doc.add_heading('Section 6: Mixed References in Parentheses', level=1)
    doc.add_paragraph('The gospels record Jesus\' ministry (mt 4:23, mk 1:14-15, lk 4:14-15).')
    doc.add_paragraph('The epistles of John emphasize love (1 jn 4:7-8, 2 jn 5-6, 3 jn 11).')
    
    # Save the document
    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    document_path = create_test_document()
    print(f"Test document created: {document_path}")

